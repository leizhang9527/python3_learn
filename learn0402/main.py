# -*- coding: utf-8 -*-
from PyQt5.QtCore import pyqtSlot
from PyQt5.QtWidgets import QMainWindow, QMessageBox, QInputDialog, QLineEdit,QDialog, QFileDialog
from PyQt5 import QtWidgets, QtGui, QtCore
from Ui_test2 import Ui_MainWindow
from Ui_info import Ui_dialog
import webbrowser
import time
import docx

def read_docx(filename):
    doc = docx.Document(filename)
    fulltext= []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return '\n'.join(fulltext)

class dialog(QDialog, Ui_dialog):
    def __init__(self, parent=None):
        super(dialog, self).__init__(parent)
        self.setupUi(self)
    
    @pyqtSlot()
    def on_pushButton_clicked(self):
       print("OK")
       my_str = self.lineEdit.text()
       my_str2 = self.lineEdit_2.text()
       my_str3 = self.lineEdit_3.text()
       my_str4 = self.lineEdit_4.text()
       my_str5 = self.lineEdit_5.text()
       my_str6 = self.lineEdit_6.text()
       print(my_str)
       print(my_str2)
       print(my_str3)
       print(my_str4)
       print(my_str5)
       print(my_str6)
       self.close()

    
    @pyqtSlot()
    def on_pushButton_2_clicked(self):
        print("cannel")
    
    @pyqtSlot()
    def on_toolButton_clicked(self):
        my_str,okPressed = QInputDialog.getInt(self, "年龄","请输入",  18, 0, 250)
        print(my_str)
    
    @pyqtSlot()
    def on_toolButton_2_clicked(self):
        list = ["男","女", "秘密"]
        my_str,okPressed = QInputDialog.getItem(self, "性别","请选择", list)
        print(my_str)
    
    @pyqtSlot()
    def on_toolButton_3_clicked(self):
        my_str,okPressed = QInputDialog.getInt(self, "身高","请输入",  180, 0, 250)
        print(my_str)
    
    @pyqtSlot()
    def on_toolButton_4_clicked(self):
    
        my_str,okPressed = QInputDialog.getInt(self, "体重","请输入",  60, 0, 100)
        print(my_str)
    

class MainWindow(QMainWindow, Ui_MainWindow):
    def __init__(self, parent=None):
        super(MainWindow, self).__init__(parent)
        self.setupUi(self)
        self.graphicsView.mousePressEvent = self.myclicked
        time.sleep(2)

    def myclicked(self, e):
        print("clicked!")
        webbrowser.open('https://www.baidu.com/')
        
    @pyqtSlot()
    def on_pushButton_5_clicked(self):
        print ('按钮被按下')
        self.label_5.setText( '卧槽，你居然敢按我！')


    @pyqtSlot()
    def on_pushButton_6_clicked(self):
      self.lineEdit.setText('')

    @pyqtSlot()
    def on_pushButton_7_clicked(self):
        my_str = self.lineEdit.text()
        self.textEdit.append(my_str)
        print(self.lineEdit.text())

    @pyqtSlot()
    def on_pushButton_8_clicked(self):
        msg = self.textEdit.toPlainText()
        print(msg)

    @pyqtSlot()
    def on_radioButton_3_clicked(self):
        print('you select r3')

    @pyqtSlot()
    def on_radioButton_clicked(self):
        print('you select r1')
        self.radioButton_4.setChecked(True)
#        if self.radioButton_4.isChecked():
#            print('r4 is checked')
#        elif self.radioButton_5.isChecked():
#            print('r5 is checked')


    @pyqtSlot()
    def on_radioButton_2_clicked(self):
        print('you select r2')
        self.radioButton_5.setChecked(True)

    @pyqtSlot()
    def on_radioButton_5_clicked(self):
        print('you select r5')

    @pyqtSlot()
    def on_radioButton_4_clicked(self):
       print('you select r4')


    @pyqtSlot(int)
    def on_dial_valueChanged(self, value):
       print(value)
       self.lcdNumber.display(value)

    @pyqtSlot(int)
    def on_horizontalSlider_valueChanged(self, value):
        print(value)
        self.lcdNumber.display(value)

    @pyqtSlot(int)
    def on_verticalSlider_valueChanged(self, value):
        print(value)
        self.lcdNumber.display(value)

    @pyqtSlot()
    def on_pushButton_9_clicked(self):
       QMessageBox.information(self,"请回答","你最帅？")

    @pyqtSlot()
    def on_pushButton_10_clicked(self):
        QMessageBox.question(self,"请回答","你最帅？",QMessageBox.Yes | QMessageBox.No)
        
    @pyqtSlot()
    def on_pushButton_11_clicked(self):
       reply = QMessageBox.warning(self,'警告','文字编码方式不同',QMessageBox.Yes | QMessageBox.No| QMessageBox.Cancel)
       print(reply)
       if reply == 16384:
           print ('保存')
       else:
            print ('不保存')
    
    @pyqtSlot()
    def on_pushButton_12_clicked(self):
       reply1 = QMessageBox.critical(self,'严重警告','文字编码方式不同',QMessageBox.Yes | QMessageBox.No)
       print(reply1)
    
    @pyqtSlot()
    def on_pushButton_13_clicked(self):
        reply = QMessageBox.about(self,'关于','这是关于一个对话框的练习')
        print(reply)
     
    @pyqtSlot()
    def on_pushButton_14_clicked(self):
         reply = QMessageBox.aboutQt(self, "介绍Qt")
         print(reply)
        
    @pyqtSlot()
    def on_pushButton_15_clicked(self):
       my_str,okPressed = QInputDialog.getText(self, "字符串","请在此输入", QLineEdit.Normal, "请在此输入信息")
       print(my_str)
       
    @pyqtSlot()
    def on_pushButton_16_clicked(self):
       my_str,okPressed = QInputDialog.getInt(self, "数字","请输入一个整数",  66, 0, 1000)
       print(my_str)
       
    @pyqtSlot()
    def on_pushButton_17_clicked(self):
       my_str,okPressed = QInputDialog.getDouble(self, "浮点数","请输入一个浮点数",  66.666, 0.6666, 1000.6666)
       print(my_str)
  
    @pyqtSlot()
    def on_pushButton_18_clicked(self):
        list = ["龙眼","苹果", "草莓"]
        my_str,okPressed = QInputDialog.getItem(self, "今天吃什么","请选择", list)
        print(my_str)
        self.textEdit.append(my_str)
    
    @pyqtSlot()
    def on_pushButton_19_clicked(self):
        my_info = dialog()
        my_info.exec()
    
    @pyqtSlot()
    def on_pushButton_20_clicked(self):
        self.graphicsView.setStyleSheet("border-image: url(:/pic/source/pic/3.jpg);")
     
    @pyqtSlot()
    def on_action_dakai_triggered(self):
        print("打开！")
        file_path , file_type= QFileDialog.getOpenFileName(self,"请选择文件","D:/work-space/eric6/learn0402/")
        print (file_path)
        print (file_type)
        if file_path[-4:] == '.doc' or  file_path[-5:] == '.docx':
            print("word")
            from win32com import client as wc
            word = wc.Dispatch('Word.Application')
            word.Visible = 0
            self.textEdit.append(read_docx(file_path.replace('/', '\\')))
            
#            my_worddoc = word.Documents.Open(file_path.replace('/', '\\'))
#            my_count = my_worddoc.Paragraphs.Count
#            for i in range(my_count):
#                my_pr = my_worddoc.Paragraphs[i].Range
#                print(my_pr.text)
#                self.textEdit.append(my_pr.text)
#            my_worddoc.Close()

#            import docx
#            doc = docx.Document(file_path.replace('/', '\\'))
#            for my_para in doc.paragraphs:
#                print(my_para.text)
#                self.textEdit.append(my_para.text)
        elif file_path[-4:] == '.txt':
            f= open(file_path, mode = 'r', encoding='gbk')
            self.textEdit.setText('')
            for line in f.readlines():
                print(line)
                self.textEdit.append(line)
            f.close()
        elif file_path.endwith('.xlsx'):
               print("excel!")
#            from win32com import client as wc
#            excel = wc.Dispatch('Excel.Application')
#            excel.Visible = 0
#            my_excel = excel.Workbooks.Open(file_path.replace('/', '\\'))
#            print(my_excel.Sheets.Count)
#            my_sheet = my_excel.Sheets('Sheet1')
#            print(my_sheet.UsedRange.Rows.Count)
#            print(my_sheet.UsedRange.Columns.Count)
#            for i in range(my_sheet.UsedRange.Rows.Count):
#                for j in range(my_sheet.UsedRange.Columns.Count):
#                    if my_sheet.Cells(i+1, j+1).Value:
#                        print(my_sheet.Cells(i+1, j+1).Value)
#            my_excel.Close()
#            excel.Quit()
           
            
           
        else:
            #QMessageBox.Information(self,"information", "不支持的文件格式")
            reply = QMessageBox.warning(self,'警告','不支持的文件格式')
            print(reply)
    
    @pyqtSlot()
    def on_action_baocun_triggered(self):
        print("保存！")
        msg = self.textEdit.toPlainText()
        my_file , file_type= QFileDialog.getSaveFileName(self,"文件另存为","D:/work-space/eric6/learn0402/")
        print (my_file)
        print (file_type)
        f= open(my_file, mode = 'w', encoding='gbk')
        f.write(msg)
        f.close()
        

    
    @pyqtSlot()
    def on_action_guanbi_triggered(self):
        print("关闭！")
        sys.exit(0)

    @pyqtSlot()
    def on_action_about_triggered(self):
        reply = QMessageBox.about(self,'关于','这是关于一个对话框的练习')
        print(reply)
        
if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    splash = QtWidgets.QSplashScreen(QtGui.QPixmap(":/pic/source/pic/1.png"))
    splash.show()
    splash.showMessage("努力加载中(:", QtCore.Qt.AlignCenter | QtCore.Qt.AlignBottom, QtCore.Qt.black)
    time.sleep(0.1)
    splash.showMessage("加载中... 0%", QtCore.Qt.AlignCenter | QtCore.Qt.AlignBottom, QtCore.Qt.black)
    time.sleep(0.1)
    splash.showMessage("加载中... 100%", QtCore.Qt.AlignCenter | QtCore.Qt.AlignBottom, QtCore.Qt.black)
    time.sleep(0.2)
    app.processEvents()
    ui = MainWindow()
    ui.show()
    splash.finish(ui)
    sys.exit(app.exec_())
    
    
